"""
Document parsers for importing chapter data from DOCX and PDF files.
Enhanced with comprehensive section parsing and validation.
"""

import json
import re
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from utils.logger import get_logger
from .models.base import ChapterData

logger = get_logger(__name__)


class ImportResult:
    """Result of an import operation with detailed feedback."""

    def __init__(self):
        self.success: bool = False
        self.data: Optional[ChapterData] = None
        self.errors: List[str] = []
        self.warnings: List[str] = []
        self.extracted_fields: Dict[str, bool] = {}
        self.section_stats: Dict[str, int] = {}

    def add_error(self, msg: str):
        self.errors.append(msg)

    def add_warning(self, msg: str):
        self.warnings.append(msg)

    def mark_field(self, field: str, extracted: bool):
        self.extracted_fields[field] = extracted

    def set_section_count(self, section: str, count: int):
        self.section_stats[section] = count

    def get_summary(self) -> str:
        """Get a human-readable summary of what was imported."""
        lines = []

        # Count extracted fields
        extracted_count = sum(1 for v in self.extracted_fields.values() if v)
        total_fields = len(self.extracted_fields)

        if extracted_count > 0:
            lines.append(f"Extracted {extracted_count}/{total_fields} fields")

        # Section stats
        for section, count in self.section_stats.items():
            if count > 0:
                lines.append(f"  - {section}: {count} items")

        if self.warnings:
            lines.append(f"Warnings: {len(self.warnings)}")

        return "\n".join(lines) if lines else "No data extracted"


class JsonValidator:
    """Validates JSON import data structure."""

    REQUIRED_FIELDS = ['chapter_data']
    CHAPTER_REQUIRED = ['class_num', 'subject', 'chapter_number']

    @classmethod
    def validate(cls, data: Dict) -> Tuple[bool, List[str]]:
        """
        Validate JSON import data structure.

        Returns:
            Tuple of (is_valid, list of error messages)
        """
        errors = []

        # Check required top-level fields
        for field in cls.REQUIRED_FIELDS:
            if field not in data:
                errors.append(f"Missing required field: '{field}'")

        if 'chapter_data' in data:
            chapter = data['chapter_data']

            # Check required chapter fields
            for field in cls.CHAPTER_REQUIRED:
                if field not in chapter:
                    errors.append(f"Missing required chapter field: '{field}'")

            # Validate field types
            if 'class_num' in chapter:
                if not isinstance(chapter['class_num'], int):
                    errors.append("'class_num' must be an integer")
                elif chapter['class_num'] not in [9, 10, 11, 12]:
                    errors.append("'class_num' must be 9, 10, 11, or 12")

            if 'chapter_number' in chapter:
                if not isinstance(chapter['chapter_number'], int):
                    errors.append("'chapter_number' must be an integer")
                elif chapter['chapter_number'] < 1:
                    errors.append("'chapter_number' must be positive")

            if 'subject' in chapter:
                valid_subjects = ['history', 'geography', 'civics', 'economics']
                if chapter['subject'] not in valid_subjects:
                    errors.append(f"'subject' must be one of: {valid_subjects}")

            # Validate list fields
            list_fields = ['concepts', 'pyq_items', 'model_answers', 'mcqs',
                           'short_answer', 'long_answer', 'map_items']
            for field in list_fields:
                if field in chapter and not isinstance(chapter[field], list):
                    errors.append(f"'{field}' must be a list")

        return len(errors) == 0, errors

    @classmethod
    def validate_json_string(cls, json_str: str) -> Tuple[bool, Dict, List[str]]:
        """
        Parse and validate a JSON string.

        Returns:
            Tuple of (is_valid, parsed_data, error messages)
        """
        try:
            data = json.loads(json_str)
        except json.JSONDecodeError as e:
            return False, {}, [f"Invalid JSON: {str(e)}"]

        is_valid, errors = cls.validate(data)
        return is_valid, data, errors


class DocxParser:
    """Parse DOCX files to extract chapter data."""

    @staticmethod
    def parse(file_bytes: bytes) -> Optional[ChapterData]:
        """
        Parse a DOCX file and extract chapter data.

        Args:
            file_bytes: The DOCX file content as bytes

        Returns:
            ChapterData object if successful, None otherwise
        """
        try:
            from docx import Document

            doc = Document(BytesIO(file_bytes))

            # Extract text from all paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            text_content = '\n'.join(full_text)

            # Parse the content
            return DocxParser._extract_chapter_data(text_content)

        except Exception as e:
            logger.warning("Failed to parse DOCX file: %s", e)
            return None

    @staticmethod
    def _extract_chapter_data(text: str) -> ChapterData:
        """Extract chapter data from text content."""
        data = ChapterData()

        # Extract class number (e.g., "CBSE Class 10")
        class_match = re.search(r'CBSE Class (\d+)', text)
        if class_match:
            data.class_num = int(class_match.group(1))

        # Extract subject
        subject_match = re.search(r'Social Science \| (\w+)', text)
        if subject_match:
            data.subject = subject_match.group(1).lower()

        # Extract chapter number (e.g., "CHAPTER 1" or "CHAPTER 2")
        chapter_match = re.search(r'CHAPTER (\d+)', text, re.IGNORECASE)
        if chapter_match:
            data.chapter_number = int(chapter_match.group(1))

        # Extract chapter title (text after CHAPTER X, before next section)
        title_match = re.search(r'CHAPTER \d+\s*\n?\s*([^\n]+)', text, re.IGNORECASE)
        if title_match:
            data.chapter_title = title_match.group(1).strip()

        # Extract weightage
        weightage_match = re.search(r'Weightage[:\s]*([^\n]+)', text)
        if weightage_match:
            data.weightage = weightage_match.group(1).strip()

        # Extract importance
        importance_match = re.search(r'Importance[:\s]*(\w+)', text)
        if importance_match:
            data.importance = importance_match.group(1).strip()

        # Extract map work
        map_match = re.search(r'Map Work[:\s]*(\w+)', text)
        if map_match:
            data.map_work = map_match.group(1).strip()

        # Extract PYQ frequency
        freq_match = re.search(r'PYQ Frequency[:\s]*([^\n]+)', text)
        if freq_match:
            data.pyq_frequency = freq_match.group(1).strip()

        # Extract learning objectives
        obj_match = re.search(r'Learning Objectives[:\s]*\n(.*?)(?=\*|Part [A-G]|$)', text, re.DOTALL | re.IGNORECASE)
        if obj_match:
            data.learning_objectives = obj_match.group(1).strip()

        # Extract syllabus alert
        alert_match = re.search(r'SYLLABUS ALERT[:\s]*([^\n]+)', text)
        if alert_match:
            data.syllabus_alert_text = alert_match.group(1).strip()
            data.syllabus_alert_enabled = True

        return data

    @staticmethod
    def parse_section(file_bytes: bytes, section: str) -> Dict[str, Any]:
        """
        Parse a specific section from a DOCX file.

        Args:
            file_bytes: The DOCX file content as bytes
            section: Section to parse ('cover', 'part_a', 'part_b', etc.)

        Returns:
            Dict with section-specific data
        """
        try:
            from docx import Document

            doc = Document(BytesIO(file_bytes))

            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            text_content = '\n'.join(full_text)

            # Route to section-specific parser
            parsers = {
                'cover': DocxParser._parse_cover_section,
                'part_a': DocxParser._parse_part_a_section,
                'part_b': DocxParser._parse_part_b_section,
                'part_c': DocxParser._parse_part_c_section,
                'part_d': DocxParser._parse_part_d_section,
                'part_e': DocxParser._parse_part_e_section,
                'part_f': DocxParser._parse_part_f_section,
                'part_g': DocxParser._parse_part_g_section,
            }

            parser_func = parsers.get(section)
            if parser_func:
                return parser_func(text_content)
            return {}

        except Exception as e:
            logger.warning("Failed to parse DOCX section '%s': %s", section, e)
            return {}

    @staticmethod
    def _parse_cover_section(text: str) -> Dict[str, Any]:
        """Parse cover page data."""
        data = {}

        # Chapter title
        title_match = re.search(r'CHAPTER \d+\s*\n?\s*([^\n]+)', text, re.IGNORECASE)
        if title_match:
            data['chapter_title'] = title_match.group(1).strip()

        # Subtitle
        subtitle_match = re.search(r'CHAPTER \d+\s*\n[^\n]+\n([^\n]+)', text)
        if subtitle_match and not subtitle_match.group(1).startswith(('-', 'Weightage', 'CBSE')):
            data['subtitle'] = subtitle_match.group(1).strip()

        # Weightage, Importance, etc.
        for field, pattern in [
            ('weightage', r'Weightage[:\s]*([^\n]+)'),
            ('importance', r'Importance[:\s]*(\w+)'),
            ('map_work', r'Map Work[:\s]*(\w+)'),
            ('pyq_frequency', r'PYQ Frequency[:\s]*([^\n]+)'),
        ]:
            match = re.search(pattern, text)
            if match:
                data[field] = match.group(1).strip()

        # Learning objectives
        obj_match = re.search(r'Learning Objectives[:\s]*\n(.*?)(?=\*|Chapter Contents|$)', text, re.DOTALL | re.IGNORECASE)
        if obj_match:
            data['learning_objectives'] = obj_match.group(1).strip()

        # Syllabus alert
        alert_match = re.search(r'SYLLABUS ALERT[:\s]*([^\n]+)', text)
        if alert_match:
            data['syllabus_alert_text'] = alert_match.group(1).strip()
            data['syllabus_alert_enabled'] = True

        return data

    @staticmethod
    def _parse_part_a_section(text: str) -> Dict[str, Any]:
        """Parse Part A: PYQ Analysis data."""
        data = {'pyq_items': []}

        # Year range
        range_match = re.search(r'PYQ.*?(\d{4}-\d{4})', text)
        if range_match:
            data['pyq_year_range'] = range_match.group(1)

        # Prediction
        pred_match = re.search(r'Prediction[:\s]*([^\n]+)', text)
        if pred_match:
            data['pyq_prediction'] = pred_match.group(1).strip()

        # Syllabus note
        note_match = re.search(r'Syllabus Note[:\s]*([^\n]+)', text)
        if note_match:
            data['pyq_syllabus_note'] = note_match.group(1).strip()

        return data

    @staticmethod
    def _parse_part_b_section(text: str) -> Dict[str, Any]:
        """Parse Part B: Key Concepts data."""
        data = {'concepts': []}

        # Try multiple patterns to find Part B section
        part_b_patterns = [
            r'Part B[:\s]*Key Concepts(.*?)(?=Part [C-G]|End of Chapter|$)',
            r'Key Concepts(.*?)(?=Part [C-G]|Model Answers|$)',
            r'Part B(.*?)(?=Part C|$)',
        ]

        part_b_text = None
        for pattern in part_b_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                part_b_text = match.group(1)
                break

        if not part_b_text:
            return data

        # Find numbered concepts - try multiple patterns
        concept_patterns = [
            r'(\d+)\.\s*([^\n]+)',  # "1. Title"
            r'Concept\s*(\d+)[:\s]*([^\n]+)',  # "Concept 1: Title"
            r'(\d+)\)\s*([^\n]+)',  # "1) Title"
        ]

        concepts = []
        for pattern in concept_patterns:
            matches = list(re.finditer(pattern, part_b_text))
            if matches:
                for match in matches:
                    num = match.group(1)
                    title = match.group(2).strip()
                    if title and len(title) > 3 and not title.startswith(('(', '-', '•', 'mark')):
                        concepts.append({
                            'number': int(num),
                            'title': title,
                            'content': '',
                            'ncert_line': '',
                            'memory_trick': '',
                            'did_you_know': ''
                        })
                break

        if concepts:
            data['concepts'] = concepts

        return data

    @staticmethod
    def _parse_part_c_section(text: str) -> Dict[str, Any]:
        """Parse Part C: Model Answers data."""
        data = {'model_answers': []}

        # Find Part C section
        part_c_match = re.search(
            r'Part C[:\s]*Model Answers(.*?)(?=Part [D-G]|End of Chapter|$)',
            text, re.DOTALL | re.IGNORECASE
        )

        if not part_c_match:
            return data

        part_c_text = part_c_match.group(1)

        # Find questions with marks - patterns like "Q1. Question text [3M]" or "1. Question [3 Marks]"
        question_patterns = [
            r'Q?(\d+)[\.\)]\s*(.+?)\s*\[(\d+)\s*[Mm](?:arks?)?\]',
            r'Q?(\d+)[\.\)]\s*(.+?)(?=\nAnswer|\nAns)',
        ]

        for pattern in question_patterns:
            matches = list(re.finditer(pattern, part_c_text, re.DOTALL))
            if matches:
                for match in matches:
                    num = match.group(1)
                    question = match.group(2).strip()
                    marks = match.group(3) if len(match.groups()) > 2 else '3'

                    # Try to find the answer following this question
                    answer = ''
                    ans_match = re.search(
                        rf'Q?{num}[\.\)].*?(?:Answer|Ans)[:\s]*(.+?)(?=Q?\d+[\.\)]|$)',
                        part_c_text, re.DOTALL | re.IGNORECASE
                    )
                    if ans_match:
                        answer = ans_match.group(1).strip()[:500]  # Limit answer length

                    if question and len(question) > 10:
                        data['model_answers'].append({
                            'question': question[:200],  # Limit question length
                            'answer': answer,
                            'marks': int(marks) if marks.isdigit() else 3,
                            'marking_points': []
                        })
                break

        return data

    @staticmethod
    def _parse_part_d_section(text: str) -> Dict[str, Any]:
        """Parse Part D: Practice Questions data."""
        data = {
            'mcqs': [],
            'short_answer': [],
            'long_answer': [],
            'assertion_reason': []
        }

        # Find Part D section
        part_d_match = re.search(
            r'Part D[:\s]*Practice Questions(.*?)(?=Part [E-G]|End of Chapter|$)',
            text, re.DOTALL | re.IGNORECASE
        )

        if not part_d_match:
            return data

        part_d_text = part_d_match.group(1)

        # Parse MCQs - look for questions with (a) (b) (c) (d) options
        mcq_pattern = r'(\d+)[\.\)]\s*(.+?)\s*\n\s*\(a\)\s*(.+?)\n\s*\(b\)\s*(.+?)\n\s*\(c\)\s*(.+?)\n\s*\(d\)\s*(.+?)(?=\n\d+[\.\)]|\n\s*(?:Answer|Short|Long|$))'
        mcq_matches = re.finditer(mcq_pattern, part_d_text, re.DOTALL | re.IGNORECASE)

        for match in mcq_matches:
            question = match.group(2).strip()
            if question and len(question) > 5:
                data['mcqs'].append({
                    'question': question[:200],
                    'options': [
                        match.group(3).strip()[:100],
                        match.group(4).strip()[:100],
                        match.group(5).strip()[:100],
                        match.group(6).strip()[:100]
                    ],
                    'answer': '',
                    'difficulty': 'M'
                })

        # Parse Short Answer Questions (3 marks)
        sa_section = re.search(
            r'Short Answer.*?(?:3 [Mm]arks?)(.*?)(?=Long Answer|5 [Mm]arks?|Part [E-G]|$)',
            part_d_text, re.DOTALL | re.IGNORECASE
        )
        if sa_section:
            sa_questions = re.findall(r'(\d+)[\.\)]\s*(.+?)(?=\n\d+[\.\)]|$)', sa_section.group(1))
            for _, q in sa_questions:
                q = q.strip()
                if q and len(q) > 10 and not q.startswith(('(', 'Answer')):
                    data['short_answer'].append({
                        'question': q[:200],
                        'marks': 3,
                        'hint': ''
                    })

        # Parse Long Answer Questions (5 marks)
        la_section = re.search(
            r'Long Answer.*?(?:5 [Mm]arks?)(.*?)(?=HOTS|Part [E-G]|$)',
            part_d_text, re.DOTALL | re.IGNORECASE
        )
        if la_section:
            la_questions = re.findall(r'(\d+)[\.\)]\s*(.+?)(?=\n\d+[\.\)]|$)', la_section.group(1))
            for _, q in la_questions:
                q = q.strip()
                if q and len(q) > 10 and not q.startswith(('(', 'Answer')):
                    data['long_answer'].append({
                        'question': q[:300],
                        'marks': 5,
                        'hint': ''
                    })

        return data

    @staticmethod
    def _parse_part_e_section(text: str) -> Dict[str, Any]:
        """Parse Part E: Map Work data."""
        data = {}

        # Map items
        items = re.findall(r'\d+\.\s*([^\n]+?)(?=\d+\.|$)', text)
        if items:
            data['map_items'] = [item.strip() for item in items if item.strip()]

        # Map tips
        tips_match = re.search(r'Map.*Tips[:\s]*\n(.*?)(?=Part|$)', text, re.DOTALL | re.IGNORECASE)
        if tips_match:
            data['map_tips'] = tips_match.group(1).strip()

        return data

    @staticmethod
    def _parse_part_f_section(text: str) -> Dict[str, Any]:
        """Parse Part F: Quick Revision data."""
        data = {
            'revision_key_points': [],
            'revision_key_terms': [],
            'revision_memory_tricks': []
        }

        # Find Part F section
        part_f_match = re.search(
            r'Part F[:\s]*Quick Revision(.*?)(?=Part G|End of Chapter|$)',
            text, re.DOTALL | re.IGNORECASE
        )

        if not part_f_match:
            return data

        part_f_text = part_f_match.group(1)

        # Parse Key Points
        key_points_section = re.search(
            r'Key Points?(.*?)(?=Key Terms?|Memory|Timeline|$)',
            part_f_text, re.DOTALL | re.IGNORECASE
        )
        if key_points_section:
            points = re.findall(r'[\d\•\-\*]\s*(.+?)(?=[\n\d\•\-\*]|$)', key_points_section.group(1))
            for point in points:
                point = point.strip()
                if point and len(point) > 10:
                    data['revision_key_points'].append(point[:200])

        # Parse Key Terms (Term: Definition format)
        terms_section = re.search(
            r'Key Terms?(.*?)(?=Memory|Timeline|Part G|$)',
            part_f_text, re.DOTALL | re.IGNORECASE
        )
        if terms_section:
            # Look for "Term: Definition" or "Term - Definition" patterns
            term_matches = re.findall(
                r'([A-Z][a-zA-Z\s]+?)[\s]*[:\-–]\s*(.+?)(?=\n[A-Z]|\n\n|$)',
                terms_section.group(1)
            )
            for term, definition in term_matches:
                term = term.strip()
                definition = definition.strip()
                if term and definition and len(term) > 2 and len(definition) > 5:
                    data['revision_key_terms'].append({
                        'term': term[:50],
                        'definition': definition[:200]
                    })

        # Parse Memory Tricks
        memory_section = re.search(
            r'Memory Tricks?(.*?)(?=Part G|End of Chapter|$)',
            part_f_text, re.DOTALL | re.IGNORECASE
        )
        if memory_section:
            tricks = re.findall(r'[\d\•\-\*>]\s*(.+?)(?=[\n\d\•\-\*>]|$)', memory_section.group(1))
            for trick in tricks:
                trick = trick.strip()
                if trick and len(trick) > 5:
                    data['revision_memory_tricks'].append(trick[:150])

        return data

    @staticmethod
    def _parse_part_g_section(text: str) -> Dict[str, Any]:
        """Parse Part G: Exam Strategy data."""
        data = {
            'time_allocation': [],
            'common_mistakes_exam': [],
            'examiner_pro_tips': [],
            'self_assessment_checklist': []
        }

        # Find Part G section
        part_g_match = re.search(
            r'Part G[:\s]*Exam Strategy(.*?)(?=End of Chapter|$)',
            text, re.DOTALL | re.IGNORECASE
        )

        if not part_g_match:
            return data

        part_g_text = part_g_match.group(1)

        # Parse Time Allocation table
        time_section = re.search(
            r'Time Allocation(.*?)(?=Common Mistakes|What Loses|Pro Tips|$)',
            part_g_text, re.DOTALL | re.IGNORECASE
        )
        if time_section:
            # Look for "Question Type | Marks | Time" or similar patterns
            time_matches = re.findall(
                r'(MCQ|Short|Long|Map|Source|Case).*?(\d+)\s*[Mm]arks?.*?(\d+)\s*[Mm]in',
                time_section.group(1), re.IGNORECASE
            )
            for q_type, marks, time in time_matches:
                data['time_allocation'].append({
                    'type': q_type.strip(),
                    'marks': marks,
                    'time': f"{time} min"
                })

        # Parse Common Mistakes
        mistakes_section = re.search(
            r'(?:Common Mistakes|What Loses Marks)(.*?)(?=Pro Tips|Examiner|Checklist|$)',
            part_g_text, re.DOTALL | re.IGNORECASE
        )
        if mistakes_section:
            # Look for "Mistake | Correction" or "X ... -> ..." patterns
            mistake_matches = re.findall(
                r'[✗❌\-]\s*(.+?)(?:→|->|:)\s*(.+?)(?=\n[✗❌\-]|\n\n|$)',
                mistakes_section.group(1)
            )
            for mistake, correction in mistake_matches:
                if mistake.strip() and correction.strip():
                    data['common_mistakes_exam'].append({
                        'mistake': mistake.strip()[:150],
                        'correction': correction.strip()[:150]
                    })

        # Parse Pro Tips
        tips_section = re.search(
            r'(?:Pro Tips|Examiner.*?Tips)(.*?)(?=Checklist|Self.*?Assessment|$)',
            part_g_text, re.DOTALL | re.IGNORECASE
        )
        if tips_section:
            tips = re.findall(r'[✓✔\d\•\-\*]\s*(.+?)(?=[\n✓✔\d\•\-\*]|$)', tips_section.group(1))
            for tip in tips:
                tip = tip.strip()
                if tip and len(tip) > 10:
                    data['examiner_pro_tips'].append(tip[:200])

        # Parse Self-Assessment Checklist
        checklist_section = re.search(
            r'(?:Self.*?Assessment|Checklist)(.*?)(?=End of|$)',
            part_g_text, re.DOTALL | re.IGNORECASE
        )
        if checklist_section:
            items = re.findall(r'[☐□\[\]\d\•\-]\s*(.+?)(?=[\n☐□\[\]\d\•\-]|$)', checklist_section.group(1))
            for item in items:
                item = item.strip()
                if item and len(item) > 5:
                    data['self_assessment_checklist'].append(item[:150])

        return data


class PdfParser:
    """Parse PDF files to extract chapter data."""

    # Track if last parse failure was due to missing dependencies
    _last_failure_missing_deps: bool = False

    @staticmethod
    def is_available() -> bool:
        """Check if PDF parsing dependencies are installed."""
        try:
            import fitz  # PyMuPDF
            return True
        except ImportError:
            pass

        try:
            import pdfplumber
            return True
        except ImportError:
            pass

        return False

    @staticmethod
    def get_missing_dependency_message() -> str:
        """Get user-friendly message about missing PDF dependencies."""
        return (
            "PDF import requires PyMuPDF or pdfplumber.\n"
            "Install with: `pip install pymupdf` or `pip install pdfplumber`"
        )

    @classmethod
    def was_missing_deps(cls) -> bool:
        """Check if the last parse failure was due to missing dependencies."""
        return cls._last_failure_missing_deps

    @classmethod
    def parse(cls, file_bytes: bytes) -> Optional[ChapterData]:
        """
        Parse a PDF file and extract chapter data.

        Args:
            file_bytes: The PDF file content as bytes

        Returns:
            ChapterData object if successful, None otherwise
        """
        # Reset the missing deps flag
        cls._last_failure_missing_deps = False

        try:
            # Try PyMuPDF first
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(stream=file_bytes, filetype="pdf")
                text_content = ""
                for page in doc:
                    text_content += page.get_text()
                doc.close()

            except ImportError:
                # Fallback to pdfplumber
                try:
                    import pdfplumber

                    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                        text_content = ""
                        for page in pdf.pages:
                            text_content += page.extract_text() or ""

                except ImportError:
                    # Neither library is available
                    cls._last_failure_missing_deps = True
                    return None

            # Use the same extraction logic as DOCX
            return DocxParser._extract_chapter_data(text_content)

        except Exception as e:
            logger.warning("Failed to parse PDF file: %s", e)
            return None


class MarkdownParser:
    """Parse Markdown files to extract chapter data."""

    @staticmethod
    def parse(file_bytes: bytes) -> Optional[ChapterData]:
        """
        Parse a Markdown file and extract chapter data.

        Args:
            file_bytes: The Markdown file content as bytes

        Returns:
            ChapterData object if successful, None otherwise
        """
        try:
            text = file_bytes.decode('utf-8')
            return MarkdownParser._extract_chapter_data(text)
        except Exception as e:
            logger.warning("Failed to parse Markdown file: %s", e)
            return None

    @staticmethod
    def _extract_chapter_data(text: str) -> ChapterData:
        """Extract chapter data from markdown content."""
        data = ChapterData()

        # Extract chapter title from # heading
        title_match = re.search(r'^#\s+(?:Chapter\s+)?(\d+)?[:\.\s]*(.+?)$', text, re.MULTILINE | re.IGNORECASE)
        if title_match:
            if title_match.group(1):
                data.chapter_number = int(title_match.group(1))
            data.chapter_title = title_match.group(2).strip()

        # Extract chapter number if not in title
        if not data.chapter_number:
            ch_match = re.search(r'Chapter[:\s]*(\d+)', text, re.IGNORECASE)
            if ch_match:
                data.chapter_number = int(ch_match.group(1))

        # Extract class number
        class_match = re.search(r'Class[:\s]*(\d+)', text, re.IGNORECASE)
        if class_match:
            data.class_num = int(class_match.group(1))

        # Extract subject from frontmatter or text
        subject_patterns = [
            r'Subject[:\s]*(History|Geography|Civics|Economics)',
            r'\b(History|Geography|Civics|Economics)\b'
        ]
        for pattern in subject_patterns:
            subject_match = re.search(pattern, text, re.IGNORECASE)
            if subject_match:
                data.subject = subject_match.group(1).lower()
                break

        # Extract metadata from YAML frontmatter if present
        frontmatter_match = re.match(r'^---\s*\n(.*?)\n---', text, re.DOTALL)
        if frontmatter_match:
            frontmatter = frontmatter_match.group(1)
            MarkdownParser._parse_frontmatter(frontmatter, data)

        # Parse sections
        MarkdownParser._parse_cover_section(text, data)
        MarkdownParser._parse_pyq_section(text, data)
        MarkdownParser._parse_concepts_section(text, data)
        MarkdownParser._parse_practice_section(text, data)
        MarkdownParser._parse_revision_section(text, data)

        return data

    @staticmethod
    def _parse_frontmatter(frontmatter: str, data: ChapterData):
        """Parse YAML-like frontmatter."""
        # Simple key: value parsing
        patterns = {
            'title': r'title[:\s]*["\']?(.+?)["\']?\s*$',
            'chapter': r'chapter[:\s]*(\d+)',
            'class': r'class[:\s]*(\d+)',
            'subject': r'subject[:\s]*["\']?(.+?)["\']?\s*$',
            'weightage': r'weightage[:\s]*["\']?(.+?)["\']?\s*$',
            'importance': r'importance[:\s]*["\']?(.+?)["\']?\s*$',
            'map_work': r'map_work[:\s]*["\']?(.+?)["\']?\s*$',
        }

        for key, pattern in patterns.items():
            match = re.search(pattern, frontmatter, re.MULTILINE | re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if key == 'title':
                    data.chapter_title = value
                elif key == 'chapter':
                    data.chapter_number = int(value)
                elif key == 'class':
                    data.class_num = int(value)
                elif key == 'subject':
                    data.subject = value.lower()
                elif key == 'weightage':
                    data.weightage = value
                elif key == 'importance':
                    data.importance = value
                elif key == 'map_work':
                    data.map_work = value

    @staticmethod
    def _parse_cover_section(text: str, data: ChapterData):
        """Parse cover page elements from markdown."""
        # Learning objectives - look for ## Learning Objectives or bullet list after it
        obj_match = re.search(
            r'##\s*Learning Objectives?\s*\n((?:[-*]\s*.+\n?)+)',
            text, re.IGNORECASE
        )
        if obj_match:
            objectives = re.findall(r'[-*]\s*(.+)', obj_match.group(1))
            data.learning_objectives = '\n'.join(objectives)

        # Syllabus alert
        alert_match = re.search(
            r'(?:>\s*\*\*?|##\s*)(?:⚠|!)?.*?Syllabus Alert[:\s]*(.+?)(?:\n\n|$)',
            text, re.IGNORECASE | re.DOTALL
        )
        if alert_match:
            data.syllabus_alert_text = alert_match.group(1).strip()
            data.syllabus_alert_enabled = True

        # Weightage, Importance, etc. from table or list
        for field, attr in [('Weightage', 'weightage'), ('Importance', 'importance'),
                            ('Map Work', 'map_work'), ('PYQ Frequency', 'pyq_frequency')]:
            match = re.search(rf'{field}[:\s|]*([^\n|]+)', text, re.IGNORECASE)
            if match and not getattr(data, attr):
                setattr(data, attr, match.group(1).strip())

    @staticmethod
    def _parse_pyq_section(text: str, data: ChapterData):
        """Parse Part A: PYQ Analysis from markdown."""
        # Find Part A or PYQ section
        part_a_match = re.search(
            r'##\s*(?:Part A|PYQ|Previous Year).*?\n(.*?)(?=##\s*Part [B-G]|$)',
            text, re.DOTALL | re.IGNORECASE
        )
        if not part_a_match:
            return

        section = part_a_match.group(1)

        # Year range
        range_match = re.search(r'(\d{4})\s*[-–]\s*(\d{4})', section)
        if range_match:
            data.pyq_year_range = f"{range_match.group(1)}-{range_match.group(2)}"

        # PYQ items from table
        # Look for | Question | Marks | Years | pattern
        table_match = re.search(r'\|[^\n]+\|\s*\n\|[-:\s|]+\|\s*\n((?:\|[^\n]+\|\s*\n?)+)', section)
        if table_match:
            rows = table_match.group(1).strip().split('\n')
            for row in rows:
                cells = [c.strip() for c in row.split('|') if c.strip()]
                if len(cells) >= 2:
                    from .models.base import PYQItem
                    item = PYQItem()
                    item.question = cells[0][:200]
                    if len(cells) >= 2:
                        marks_match = re.search(r'(\d+)', cells[1])
                        if marks_match:
                            item.marks = marks_match.group(1)
                    if len(cells) >= 3:
                        item.years = cells[2]
                    if item.question:
                        data.pyq_items.append(item)

        # Prediction
        pred_match = re.search(r'(?:Prediction|🎯)[:\s]*(.+?)(?:\n|$)', section, re.IGNORECASE)
        if pred_match:
            data.pyq_prediction = pred_match.group(1).strip()

        # Syllabus note
        note_match = re.search(r'Syllabus Note[:\s]*(.+?)(?:\n\n|$)', section, re.IGNORECASE)
        if note_match:
            data.pyq_syllabus_note = note_match.group(1).strip()

    @staticmethod
    def _parse_concepts_section(text: str, data: ChapterData):
        """Parse Part B: Key Concepts from markdown."""
        # Find Part B or Concepts section
        part_b_match = re.search(
            r'##\s*(?:Part B|Key Concepts).*?\n(.*?)(?=##\s*Part [C-G]|$)',
            text, re.DOTALL | re.IGNORECASE
        )
        if not part_b_match:
            return

        section = part_b_match.group(1)

        # Find concepts by ### headers or numbered items
        concept_headers = re.findall(
            r'###\s*(\d+)?\.?\s*(.+?)\n(.*?)(?=###|\Z)',
            section, re.DOTALL
        )

        from .models.base import ConceptItem

        for i, (num, title, content) in enumerate(concept_headers, 1):
            concept = ConceptItem()
            concept.number = int(num) if num else i
            concept.title = title.strip()

            # Extract content (remove sub-sections)
            content_lines = []
            for line in content.split('\n'):
                line = line.strip()
                if line and not line.startswith(('#', '>', '**NCERT', '**Memory', '**Do You Know')):
                    # Remove markdown formatting
                    line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                    line = re.sub(r'\*(.+?)\*', r'\1', line)
                    content_lines.append(line)
            concept.content = '\n'.join(content_lines[:20])  # Limit lines

            # NCERT line
            ncert_match = re.search(r'(?:NCERT|📌)[:\s]*["\']?(.+?)["\']?(?:\n|$)', content, re.IGNORECASE)
            if ncert_match:
                concept.ncert_line = ncert_match.group(1).strip()

            # Memory trick
            memory_match = re.search(r'(?:Memory Trick|💡)[:\s]*(.+?)(?:\n|$)', content, re.IGNORECASE)
            if memory_match:
                concept.memory_trick = memory_match.group(1).strip()

            # Did you know
            dyk_match = re.search(r'(?:Do You Know|Did You Know)[:\s?]*(.+?)(?:\n\n|$)', content, re.IGNORECASE | re.DOTALL)
            if dyk_match:
                concept.did_you_know = dyk_match.group(1).strip()[:200]

            if concept.title:
                data.concepts.append(concept)

        # Common mistakes
        mistakes_match = re.search(r'(?:Common Mistakes|⚠)[:\s]*\n((?:[-*]\s*.+\n?)+)', section, re.IGNORECASE)
        if mistakes_match:
            mistakes = re.findall(r'[-*]\s*(.+)', mistakes_match.group(1))
            data.common_mistakes = [m.strip() for m in mistakes if m.strip()]

        # Important dates from table
        dates_match = re.search(r'(?:Important Dates|Timeline).*?\n\|[^\n]+\|\s*\n\|[-:\s|]+\|\s*\n((?:\|[^\n]+\|\s*\n?)+)', section, re.IGNORECASE)
        if dates_match:
            rows = dates_match.group(1).strip().split('\n')
            for row in rows:
                cells = [c.strip() for c in row.split('|') if c.strip()]
                if len(cells) >= 2:
                    data.important_dates.append({
                        'year': cells[0],
                        'event': cells[1]
                    })

    @staticmethod
    def _parse_practice_section(text: str, data: ChapterData):
        """Parse Part D: Practice Questions from markdown."""
        # Find Part D or Practice section
        part_d_match = re.search(
            r'##\s*(?:Part D|Practice Questions).*?\n(.*?)(?=##\s*Part [E-G]|$)',
            text, re.DOTALL | re.IGNORECASE
        )
        if not part_d_match:
            return

        section = part_d_match.group(1)

        from .models.base import QuestionItem

        # Parse MCQs
        mcq_section = re.search(r'(?:MCQ|Multiple Choice).*?\n(.*?)(?=###|Short Answer|Long Answer|$)',
                                 section, re.DOTALL | re.IGNORECASE)
        if mcq_section:
            # Pattern: 1. Question\n   a) opt1\n   b) opt2\n   c) opt3\n   d) opt4
            mcq_pattern = r'(\d+)\.\s*(.+?)\n\s*(?:a\)|A\.)\s*(.+?)\n\s*(?:b\)|B\.)\s*(.+?)\n\s*(?:c\)|C\.)\s*(.+?)\n\s*(?:d\)|D\.)\s*(.+?)(?=\n\d+\.|$)'
            for match in re.finditer(mcq_pattern, mcq_section.group(1), re.DOTALL):
                q = QuestionItem()
                q.question = match.group(2).strip()[:200]
                q.options = [
                    match.group(3).strip()[:100],
                    match.group(4).strip()[:100],
                    match.group(5).strip()[:100],
                    match.group(6).strip()[:100]
                ]
                q.difficulty = 'M'
                if q.question:
                    data.mcqs.append(q)

        # Parse Short Answer
        sa_section = re.search(r'(?:Short Answer|3 Marks?).*?\n(.*?)(?=###|Long Answer|5 Marks?|$)',
                                section, re.DOTALL | re.IGNORECASE)
        if sa_section:
            questions = re.findall(r'(\d+)\.\s*(.+?)(?=\n\d+\.|$)', sa_section.group(1), re.DOTALL)
            for _, q_text in questions:
                q = QuestionItem()
                q.question = q_text.strip()[:200]
                q.marks = 3
                if q.question and len(q.question) > 10:
                    data.short_answer.append(q)

        # Parse Long Answer
        la_section = re.search(r'(?:Long Answer|5 Marks?).*?\n(.*?)(?=###|HOTS|$)',
                                section, re.DOTALL | re.IGNORECASE)
        if la_section:
            questions = re.findall(r'(\d+)\.\s*(.+?)(?=\n\d+\.|$)', la_section.group(1), re.DOTALL)
            for _, q_text in questions:
                q = QuestionItem()
                q.question = q_text.strip()[:300]
                q.marks = 5
                if q.question and len(q.question) > 10:
                    data.long_answer.append(q)

    @staticmethod
    def _parse_revision_section(text: str, data: ChapterData):
        """Parse Part F: Quick Revision from markdown."""
        # Find Part F or Revision section
        part_f_match = re.search(
            r'##\s*(?:Part F|Quick Revision).*?\n(.*?)(?=##\s*Part G|$)',
            text, re.DOTALL | re.IGNORECASE
        )
        if not part_f_match:
            return

        section = part_f_match.group(1)

        # Key points
        points_match = re.search(r'(?:Key Points?).*?\n((?:[-*\d]+\.\s*.+\n?)+)', section, re.IGNORECASE)
        if points_match:
            points = re.findall(r'[-*]|\d+\.\s*(.+)', points_match.group(1))
            data.revision_key_points = [p.strip() for p in points if p and p.strip()]

        # Key terms from table
        terms_match = re.search(r'(?:Key Terms?).*?\n\|[^\n]+\|\s*\n\|[-:\s|]+\|\s*\n((?:\|[^\n]+\|\s*\n?)+)', section, re.IGNORECASE)
        if terms_match:
            rows = terms_match.group(1).strip().split('\n')
            for row in rows:
                cells = [c.strip() for c in row.split('|') if c.strip()]
                if len(cells) >= 2:
                    data.revision_key_terms.append({
                        'term': cells[0][:50],
                        'definition': cells[1][:200]
                    })

        # Memory tricks
        tricks_match = re.search(r'(?:Memory Tricks?).*?\n((?:[-*>]\s*.+\n?)+)', section, re.IGNORECASE)
        if tricks_match:
            tricks = re.findall(r'[-*>]\s*(.+)', tricks_match.group(1))
            data.revision_memory_tricks = [t.strip() for t in tricks if t.strip()]

    @staticmethod
    def get_import_summary(data: ChapterData) -> Dict[str, Any]:
        """Get summary of what was extracted from markdown."""
        return {
            'chapter_title': data.chapter_title or 'Untitled',
            'chapter_number': data.chapter_number,
            'subject': data.subject,
            'class_num': data.class_num,
            'concepts_count': len(data.concepts),
            'pyq_count': len(data.pyq_items),
            'mcq_count': len(data.mcqs),
            'short_answer_count': len(data.short_answer),
            'long_answer_count': len(data.long_answer),
            'key_points_count': len(data.revision_key_points),
            'key_terms_count': len(data.revision_key_terms),
        }


def parse_document(file_bytes: bytes, file_type: str) -> Optional[ChapterData]:
    """
    Parse a document file and extract chapter data.

    Args:
        file_bytes: The file content as bytes
        file_type: File extension ('json', 'docx', 'pdf', 'md')

    Returns:
        ChapterData object if successful, None otherwise
    """
    if file_type == 'docx':
        return DocxParser.parse(file_bytes)
    elif file_type == 'pdf':
        return PdfParser.parse(file_bytes)
    elif file_type in ('md', 'markdown'):
        return MarkdownParser.parse(file_bytes)
    else:
        return None


def parse_section(file_bytes: bytes, file_type: str, section: str) -> Dict[str, Any]:
    """
    Parse a specific section from a document.

    Args:
        file_bytes: The file content as bytes
        file_type: File extension ('docx', 'pdf')
        section: Section to parse ('cover', 'part_a', 'part_b', etc.)

    Returns:
        Dict with section-specific data
    """
    if file_type == 'docx':
        return DocxParser.parse_section(file_bytes, section)
    elif file_type == 'pdf':
        # For PDF, extract text first then use same logic
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()

            # Route to appropriate parser
            parsers = {
                'cover': DocxParser._parse_cover_section,
                'part_a': DocxParser._parse_part_a_section,
                'part_b': DocxParser._parse_part_b_section,
                'part_c': DocxParser._parse_part_c_section,
                'part_d': DocxParser._parse_part_d_section,
                'part_e': DocxParser._parse_part_e_section,
                'part_f': DocxParser._parse_part_f_section,
                'part_g': DocxParser._parse_part_g_section,
            }
            parser_func = parsers.get(section)
            return parser_func(text) if parser_func else {}
        except Exception as e:
            logger.warning("Failed to parse PDF section '%s': %s", section, e)
            return {}
    return {}
