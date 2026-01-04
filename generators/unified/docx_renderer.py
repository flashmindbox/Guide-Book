"""
DOCX Renderer for Guide Book Generator.
Renders document elements to DOCX using python-docx.
"""

from typing import List
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .content_builder import Element, ElementType, BoxType
from styles.theme import Colors, Fonts
from generators.docx.helpers import DocxHelpers
from generators.docx.styles import create_styled_document


class DocxRenderer:
    """Renders document elements to DOCX."""

    def __init__(self, page_size: str = "A4"):
        self.page_size = page_size
        self.document = None
        self.styles = None

    def render(self, elements: List[Element]) -> Document:
        """Render elements to DOCX document."""
        self.document, self.styles = create_styled_document(self.page_size)

        for element in elements:
            self._render_element(element)

        return self.document

    def render_to_bytes(self, elements: List[Element]) -> bytes:
        """Render elements to DOCX and return as bytes."""
        doc = self.render(elements)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _render_element(self, element: Element):
        """Render a single element to DOCX."""
        renderers = {
            ElementType.HEADER: self._render_header,
            ElementType.TITLE: self._render_title,
            ElementType.DECORATIVE_LINE: self._render_decorative_line,
            ElementType.METADATA_TABLE: self._render_metadata_table,
            ElementType.ALERT_BOX: self._render_alert_box,
            ElementType.INFO_BOX: self._render_info_box,
            ElementType.TIP_BOX: self._render_tip_box,
            ElementType.NEUTRAL_BOX: self._render_neutral_box,
            ElementType.PARAGRAPH: self._render_paragraph,
            ElementType.BULLET_LIST: self._render_bullet_list,
            ElementType.NUMBERED_LIST: self._render_numbered_list,
            ElementType.TABLE: self._render_table,
            ElementType.PART_HEADER: self._render_part_header,
            ElementType.SECTION_TITLE: self._render_section_title,
            ElementType.CONCEPT_BOX: self._render_concept_box,
            ElementType.QUESTION_BLOCK: self._render_question_block,
            ElementType.MCQ_BLOCK: self._render_mcq_block,
            ElementType.ANSWER_BLOCK: self._render_answer_block,
            ElementType.TIMELINE: self._render_timeline,
            ElementType.QR_CODES: self._render_qr_codes,
            ElementType.PAGE_BREAK: self._render_page_break,
            ElementType.DIVIDER: self._render_divider,
            ElementType.END_MARKER: self._render_end_marker,
        }

        renderer = renderers.get(element.type)
        if renderer:
            renderer(element)

    def _render_header(self, element: Element):
        """Render header text."""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(element.content)
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_title(self, element: Element):
        """Render chapter title section."""
        content = element.content

        # Chapter number
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        run = para.add_run(content["chapter_num"])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_PART_HEADER
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # Chapter title
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(content["title"])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_CHAPTER_TITLE
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # Thin underline
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(3)
        run = para.add_run('─' * 40)
        run.font.size = Pt(10)
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # Subtitle
        if content.get("subtitle"):
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(content["subtitle"])
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.italic = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_decorative_line(self, element: Element):
        """Render decorative line."""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run('─' * 50)
        run.font.size = Pt(11)
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

    def _render_metadata_table(self, element: Element):
        """Render metadata table."""
        content = element.content
        self.document.add_paragraph()  # Spacing

        metadata = {}
        for key, data in content.items():
            highlight_color = Colors.ACCENT_RED if data.get('highlight') else Colors.BODY_TEXT
            metadata[key] = (data['value'], highlight_color)

        DocxHelpers.create_metadata_table(self.document, metadata, self.document.styles)
        self.document.add_paragraph()  # Spacing

    def _render_box(self, title: str, text: str, bg_color: str, border_color: str, title_color: str, icon: str = ""):
        """Helper to render a styled box."""
        table = self.document.add_table(rows=1, cols=1)
        table.alignment = 1
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        DocxHelpers.set_cell_background(cell, bg_color)
        DocxHelpers.set_cell_left_border_only(cell, border_color)
        DocxHelpers.set_cell_padding(cell, 150)

        # Title
        para = cell.paragraphs[0]
        run = para.add_run(f"{icon} {title}: " if icon else f"{title}: ")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(title_color)

        # Text
        run = para.add_run(text)
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        self.document.add_paragraph()  # Spacing

    def _render_alert_box(self, element: Element):
        """Render warning/alert box."""
        content = element.content
        self._render_box(
            title=content["title"],
            text=content["text"],
            bg_color=Colors.BG_WARNING,
            border_color=Colors.BORDER_WARNING,
            title_color=Colors.ACCENT_RED,
            icon="⚠️"
        )

    def _render_info_box(self, element: Element):
        """Render info box with optional list items."""
        content = element.content

        table = self.document.add_table(rows=1, cols=1)
        table.alignment = 1
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        DocxHelpers.set_cell_background(cell, Colors.BG_INFO)
        DocxHelpers.set_cell_left_border_only(cell, Colors.BORDER_INFO)
        DocxHelpers.set_cell_padding(cell, 150)

        # Title
        para = cell.paragraphs[0]
        run = para.add_run(f"🎯 {content['title']}")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # Subtitle
        if content.get("subtitle"):
            para = cell.add_paragraph()
            run = para.add_run(content["subtitle"])
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.italic = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # Text content
        if content.get("text"):
            para = cell.add_paragraph()
            run = para.add_run(content["text"])
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # List items
        if content.get("items"):
            for item in content["items"]:
                para = cell.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.15)
                run = para.add_run("• ")
                run.font.name = Fonts.PRIMARY
                DocxHelpers.add_formatted_text(para, item)

        self.document.add_paragraph()  # Spacing

    def _render_tip_box(self, element: Element):
        """Render tip box."""
        content = element.content
        self._render_box(
            title=content["title"],
            text=content["text"],
            bg_color=Colors.BG_TIP,
            border_color=Colors.BORDER_TIP,
            title_color=Colors.SUCCESS_GREEN,
            icon="💡"
        )

    def _render_neutral_box(self, element: Element):
        """Render neutral box with items."""
        content = element.content

        table = self.document.add_table(rows=1, cols=1)
        table.alignment = 1
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        DocxHelpers.set_cell_background(cell, Colors.BG_NEUTRAL)
        DocxHelpers.set_cell_borders(cell, Colors.BORDER_NEUTRAL)
        DocxHelpers.set_cell_padding(cell, 150)

        # Title
        para = cell.paragraphs[0]
        run = para.add_run(f"📚 {content['title']}")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # Items
        if content.get("items"):
            for item in content["items"]:
                para = cell.add_paragraph()
                run = para.add_run(f"{item['label']}: ")
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY
                run.font.bold = True
                run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

                run = para.add_run(item['description'])
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY
                run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_paragraph(self, element: Element):
        """Render simple paragraph."""
        para = self.document.add_paragraph()
        run = para.add_run(element.content)
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_bullet_list(self, element: Element):
        """Render bullet list."""
        for item in element.content:
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run("• " + item)
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_numbered_list(self, element: Element):
        """Render numbered list."""
        for i, item in enumerate(element.content, 1):
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(f"{i}. {item}")
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_table(self, element: Element):
        """Render data table."""
        content = element.content
        headers = content['headers']
        rows = content['rows']

        table = self.document.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = 1

        # Header row
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            DocxHelpers.set_cell_background(cell, Colors.TABLE_HEADER_BG)
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.bold = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # Data rows
        for row_idx, row in enumerate(rows):
            for col_idx, header in enumerate(headers):
                cell = table.cell(row_idx + 1, col_idx)
                para = cell.paragraphs[0]

                if isinstance(row, dict):
                    key_lower = header.lower().replace(' ', '_')
                    value = str(row.get(key_lower, row.get(header, '')))
                else:
                    value = str(row[col_idx]) if col_idx < len(row) else ''

                run = para.add_run(value)
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY_SMALL

                # Highlight marks column
                if 'mark' in header.lower():
                    run.font.color.rgb = Colors.hex_to_rgb(Colors.ACCENT_RED)
                    run.font.bold = True
                else:
                    run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        self.document.add_paragraph()  # Spacing

    def _render_part_header(self, element: Element):
        """Render part header."""
        content = element.content

        # Add page break before part
        self.document.add_page_break()

        table = self.document.add_table(rows=1, cols=1)
        table.alignment = 1
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        DocxHelpers.set_cell_background(cell, Colors.BG_NEUTRAL)
        DocxHelpers.set_cell_borders(cell, Colors.BORDER_NEUTRAL)
        DocxHelpers.set_cell_padding(cell, 150)

        para = cell.paragraphs[0]
        run = para.add_run(f"Part {content['id']}: {content['name']}")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_PART_HEADER
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        self.document.add_paragraph()  # Spacing

    def _render_section_title(self, element: Element):
        """Render section title."""
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        run = para.add_run(element.content)
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_SECTION_TITLE
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

    def _render_concept_box(self, element: Element):
        """Render concept box with all its components."""
        content = element.content

        table = self.document.add_table(rows=1, cols=1)
        table.alignment = 1
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        DocxHelpers.set_cell_background(cell, Colors.BG_NEUTRAL)
        DocxHelpers.set_cell_borders(cell, Colors.BORDER_NEUTRAL)
        DocxHelpers.set_cell_padding(cell, 150)

        # Title
        para = cell.paragraphs[0]
        run = para.add_run(f"{content['number']}. {content['title']}")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_SECTION_TITLE
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # Content
        para = cell.add_paragraph()
        run = para.add_run(content['content'])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # NCERT line
        if content.get('ncert_line'):
            para = cell.add_paragraph()
            run = para.add_run(f'📖 NCERT: "{content["ncert_line"]}"')
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.italic = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # Memory trick
        if content.get('memory_trick'):
            para = cell.add_paragraph()
            run = para.add_run("💡 Memory Trick: ")
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.bold = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.SUCCESS_GREEN)
            run = para.add_run(content['memory_trick'])
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # Did you know
        if content.get('did_you_know'):
            para = cell.add_paragraph()
            run = para.add_run("🤔 Did You Know? ")
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.bold = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.WARNING_ORANGE)
            run = para.add_run(content['did_you_know'])
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        self.document.add_paragraph()  # Spacing

    def _render_question_block(self, element: Element):
        """Render question block."""
        content = element.content

        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.1)

        run = para.add_run(f"Q{content['number']}. ")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        run = para.add_run(content['question'])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        run = para.add_run(f" [{content['marks']}M]")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.ACCENT_RED)

        if content.get('hint'):
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(f"Hint: {content['hint']}")
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY_SMALL
            run.font.italic = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.DARK_GRAY)

    def _render_mcq_block(self, element: Element):
        """Render MCQ block with options."""
        content = element.content

        # Question
        para = self.document.add_paragraph()
        run = para.add_run(f"Q{content['number']}. ")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        run = para.add_run(content['question'])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        run = para.add_run(f" [{content['marks']}M]")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.ACCENT_RED)

        # Options
        if content.get('options'):
            for i, option in enumerate(content['options']):
                para = self.document.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.35)
                letter = chr(97 + i)  # a, b, c, d
                run = para.add_run(f"({letter}) {option}")
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY
                run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_answer_block(self, element: Element):
        """Render answer block with marking points."""
        content = element.content

        # Question
        para = self.document.add_paragraph()
        run = para.add_run(f"Q{content['number']}. ")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        run = para.add_run(content['question'])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        run = para.add_run(f" [{content['marks']}M]")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.ACCENT_RED)

        # Answer
        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        run = para.add_run("Answer: ")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.SUCCESS_GREEN)

        run = para.add_run(content['answer'])
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        # Marking points
        if content.get('marking_points'):
            for point in content['marking_points']:
                para = self.document.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.35)
                run = para.add_run("✓ ")
                run.font.name = Fonts.PRIMARY
                run.font.color.rgb = Colors.hex_to_rgb(Colors.SUCCESS_GREEN)
                run = para.add_run(point)
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY_SMALL
                run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

        self.document.add_paragraph()  # Spacing

    def _render_timeline(self, element: Element):
        """Render timeline."""
        for item in element.content:
            year = item.get('year', '')
            event = item.get('event', '')

            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.15)

            run = para.add_run(f"[{year}] ")
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.bold = True
            run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

            run = para.add_run(event)
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_BODY
            run.font.color.rgb = Colors.hex_to_rgb(Colors.BODY_TEXT)

    def _render_qr_codes(self, element: Element):
        """Render QR codes section."""
        content = element.content
        self.document.add_paragraph()

        table = self.document.add_table(rows=1, cols=1)
        table.alignment = 1
        table.columns[0].width = Inches(6.5)

        cell = table.cell(0, 0)
        DocxHelpers.set_cell_background(cell, Colors.BG_LIGHT_BLUE)
        DocxHelpers.set_cell_borders(cell, Colors.PRIMARY_BLUE)
        DocxHelpers.set_cell_padding(cell, 150)

        # Title
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("📱 Scan QR Codes to Download Practice Materials")
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)

        # QR codes
        try:
            import qrcode
            from io import BytesIO

            qr_table = cell.add_table(rows=2, cols=3)
            qr_table.alignment = 1
            qr_table.columns[0].width = Inches(2.0)
            qr_table.columns[1].width = Inches(2.5)
            qr_table.columns[2].width = Inches(2.0)

            if content.get("practice_url"):
                qr = qrcode.QRCode(version=1, box_size=6, border=2)
                qr.add_data(content["practice_url"])
                qr.make(fit=True)
                img = qr.make_image(fill_color="black", back_color="white")

                buf = BytesIO()
                img.save(buf, format='PNG')
                buf.seek(0)

                qr_cell = qr_table.cell(0, 0)
                qr_para = qr_cell.paragraphs[0]
                qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = qr_para.add_run()
                run.add_picture(buf, width=Inches(1.3))

                label_cell = qr_table.cell(1, 0)
                label_para = label_cell.paragraphs[0]
                label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = label_para.add_run("Practice Questions")
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY_SMALL
                run.font.bold = True
                run.font.color.rgb = Colors.hex_to_rgb(Colors.DARK_GRAY)

            if content.get("answers_url"):
                qr = qrcode.QRCode(version=1, box_size=6, border=2)
                qr.add_data(content["answers_url"])
                qr.make(fit=True)
                img = qr.make_image(fill_color="black", back_color="white")

                buf = BytesIO()
                img.save(buf, format='PNG')
                buf.seek(0)

                qr_cell = qr_table.cell(0, 2)
                qr_para = qr_cell.paragraphs[0]
                qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = qr_para.add_run()
                run.add_picture(buf, width=Inches(1.3))

                label_cell = qr_table.cell(1, 2)
                label_para = label_cell.paragraphs[0]
                label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = label_para.add_run("With Answers")
                run.font.name = Fonts.PRIMARY
                run.font.size = Fonts.SIZE_BODY_SMALL
                run.font.bold = True
                run.font.color.rgb = Colors.hex_to_rgb(Colors.DARK_GRAY)

        except ImportError:
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("[QR codes require 'qrcode' library]")
            run.font.name = Fonts.PRIMARY
            run.font.size = Fonts.SIZE_CAPTION
            run.font.italic = True

    def _render_page_break(self, element: Element):
        """Render page break."""
        # Already handled in part_header
        pass

    def _render_divider(self, element: Element):
        """Render divider line."""
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run('─' * 60)
        run.font.size = Pt(10)
        run.font.color.rgb = Colors.hex_to_rgb(Colors.BORDER_NEUTRAL)

    def _render_end_marker(self, element: Element):
        """Render end marker."""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(20)
        run = para.add_run(element.content)
        run.font.name = Fonts.PRIMARY
        run.font.size = Fonts.SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = Colors.hex_to_rgb(Colors.PRIMARY_BLUE)
